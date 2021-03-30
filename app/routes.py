from app import app
from flask import render_template, flash
from app.forms import NDAForm
from docx import Document

@app.route('/')
@app.route('/index')
def index():
	return render_template('index.html',title = "Home")

@app.route('/nda', methods = ['GET','POST'])
def nda():
	form = NDAForm()
	if form.validate_on_submit():
		document = Document()
		paragraph = document.add_paragraph()
		paragraph.add_run(text='Test para')
		flash ('NDA template created for {}'.format(form.partyname.data))
		document.save('testNDA.docx')
		link = ('Your download is ready')
	return render_template('NDA.html',title="NDA", form=form)